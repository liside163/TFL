# -*- coding: utf-8 -*-
"""
周报生成脚本 - 按照模板格式生成，每项任务添加时间
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

def set_chinese_font(run, font_name='宋体', font_size=11):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_weekly_report():
    """根据模板格式创建周报文档"""
    
    # 读取模板文件
    template_path = r'C:\Users\LSD\Desktop\汇报\周报\李思德小组周报0104-0110.docx'
    doc = Document(template_path)
    
    # 获取表格
    table = doc.tables[0]
    
    # ===================== 填写内容 =====================
    
    # 李思德的行（行1）
    row = table.rows[1]
    
    # 列1: 上周安排
    row.cells[1].paragraphs[0].clear()
    content_plan = """深度Optuna调优框架开发（周一-周三）
单工况迁移学习实验（周三-周五）
类别不平衡问题诊断与修复（周五-周六）
源域崩塌问题分析（周六）
总结工作内容和周报整理（周日）"""
    run = row.cells[1].paragraphs[0].add_run(content_plan)
    set_chinese_font(run, '宋体', 10)
    
    # 列2: 实际完成情况
    row.cells[2].paragraphs[0].clear()
    content_done = """完成深度Optuna调优框架开发（周一-周三）
- 扩展搜索空间包括CNN/LSTM层数、注意力机制、残差连接（周一）
- LSTM隐藏层从64扩展至192，应用warmup_cosine调度器（周二）
- 目标域准确率从基线16%提升至34%（周三）

完成6种飞行工况独立训练框架（周三-周五）
- hover工况: 59.46%准确率（周四）
- velocity工况: 58.50%准确率，优化后提升43%（周四）
- waypoint工况: 47.61%（周五）
- circling工况: 52.97%（周五）
- acce工况: 45.30%（周五）
- 平均准确率达到52.77%（周五）

类别不平衡问题修复方案实现（周五-周六）
- 诊断出Motor类占比55%导致模型退化（周五）
- 实现方案A: 简化模型架构（周六上午）
- 实现方案B: 加权采样方法（周六上午）
- 实现方案C: Focal Loss替代类别权重（周六下午）
- 添加域对齐分析输出功能，打印源/目标域准确率（周六下午）

源域崩塌问题诊断（周六）
- 发现训练集准确率仅4%的严重问题（周六下午）
- 诊断出加权采样+类别权重造成双重惩罚（周六晚）"""
    run = row.cells[2].paragraphs[0].add_run(content_done)
    set_chinese_font(run, '宋体', 10)
    
    # 列3: 下周计划
    row.cells[3].paragraphs[0].clear()
    content_next = """验证Focal Loss方案效果（周一-周二）
实现时间序列数据增强方法（周二-周三）
评估传感器特征工程影响（周三-周四）
调研MMD/CORAL/ADDA域适应方法（周四-周五）
整理实验结果和周报（周六）"""
    run = row.cells[3].paragraphs[0].add_run(content_next)
    set_chinese_font(run, '宋体', 10)
    
    # 列4: 备注
    row.cells[4].paragraphs[0].clear()
    content_note = """新增代码文件：
- models/dann_deep.py（动态DANN模型）
- optuna_tune_v2.py（深度超参优化）
- train_single_condition.py（单工况训练）
- run_all_conditions.py（批量运行）
- optuna_tune_single_condition.py（结构搜索扩展）
- scripts/sync_single_condition_params.py（参数同步）"""
    run = row.cells[4].paragraphs[0].add_run(content_note)
    set_chinese_font(run, '宋体', 10)
    
    # ===================== 保存文档 =====================
    output_path = r'C:\Users\LSD\Desktop\汇报\周报\lsd_20260111_周报.docx'
    doc.save(output_path)
    print(f'周报已生成: {output_path}')
    
    return output_path

if __name__ == '__main__':
    create_weekly_report()
