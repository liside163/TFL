#!/usr/bin/env python3
"""
周报生成脚本 - 基于项目实验记录自动生成周报Word文档
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

def set_chinese_font(paragraph, font_name='宋体', font_size=12):
    """设置中文字体"""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_weekly_report():
    """创建周报Word文档"""
    
    # 创建文档
    doc = Document()
    
    # 标题
    title = doc.add_heading('李思德小组周报', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 日期范围
    date_para = doc.add_paragraph('2026年1月4日 - 2026年1月10日')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 一、本周工作内容
    doc.add_heading('一、本周工作内容', level=1)
    
    work_content = """
    本周主要围绕基于迁移学习的无人机故障诊断研究开展工作，具体包括以下内容：

    1. 研究方案设计与实验框架搭建
       • 完成了基于DANN（Domain-Adversarial Neural Network）的域对抗迁移学习方案设计
       • 明确了HIL仿真数据→Real真实飞行数据的迁移学习目标
       • 搭建了完整的实验代码框架，包含配置管理、数据预处理、模型训练、评估等模块

    2. 数据预处理与分析
       • 实现了HIL故障状态筛选，仅保留fault_state=1的故障样本，提升数据质量
       • 针对Real数据缺少fault_state列的问题，实现了变点检测方法（change_point_detection.py）
       • 完成了11分类到7分类的调整，移除了Real域缺失的4种故障类型（Propeller, Low_Voltage, Wind_Affect, Load_Lose）

    3. 基线实验与结果分析
       • 完成了6种飞行工况（hover, waypoint, velocity, circling, acce, dece）的Source-Only基线实验
       • 基线结果：源域准确率约78%，目标域准确率约16%，验证了域偏移问题的存在

    4. Optuna超参数优化
       • 开发了基础版本optuna_tune.py和深度版本optuna_tune_v2.py
       • 通过超参数优化，目标域准确率从16%提升至34%
       • 实现了单工况独立调优脚本optuna_tune_single_condition.py

    5. 单工况迁移学习实验
       • 针对不同飞行工况特性差异大的问题，实现了单工况独立训练策略
       • 各工况最佳结果：hover(59.46%), velocity(58.50%), circling(52.97%)
       • 优化后velocity工况从14%提升至58%，提升幅度达43%

    6. 问题诊断与解决方案探索
       • 诊断出类别不平衡问题（Motor类占55%）导致模型退化
       • 发现加权采样+类别权重形成双重惩罚导致源域崩塌
       • 尝试了Focal Loss、深度架构增强（TemporalAttention+ResBlock）等方法
    """
    
    # 分段添加工作内容
    for line in work_content.strip().split('\n'):
        if line.strip():
            p = doc.add_paragraph(line.strip())
    
    doc.add_paragraph()
    
    # 二、关键实验数据
    doc.add_heading('二、关键实验数据', level=1)
    
    # 创建表格
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['工况', '源域准确率', '目标域准确率', '备注']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # 数据行
    data = [
        ['hover', '80.00%', '59.46%', '最佳工况'],
        ['waypoint', '53.62%', '47.61%', '-'],
        ['velocity', '77.27%', '58.50%', '优化后+43%'],
        ['circling', '41.18%', '52.97%', '-'],
        ['acce', '73.53%', '45.30%', '-'],
    ]
    
    for i, row_data in enumerate(data):
        for j, cell_data in enumerate(row_data):
            table.rows[i+1].cells[j].text = cell_data
    
    doc.add_paragraph()
    
    # 三、下周工作计划
    doc.add_heading('三、下周工作计划', level=1)
    
    next_week = """
    1. 继续优化各工况的Optuna超参数调优
    2. 尝试数据增强方法提升模型泛化能力
    3. 探索MMD/CORAL等其他域适应方法
    4. 完善实验结果可视化（t-SNE、混淆矩阵等）
    5. 开始撰写论文相关部分
    """
    
    for line in next_week.strip().split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    
    doc.add_paragraph()
    
    # 四、问题与困难
    doc.add_heading('四、问题与困难', level=1)
    
    problems = """
    1. 类别不平衡问题严重：Motor类样本占比过高（55%），导致模型倾向于预测多数类
    2. 负迁移现象：GRL启动后部分工况出现准确率暴跌
    3. 模型过拟合风险：复杂架构在小数据集上容易过拟合
    """
    
    for line in problems.strip().split('\n'):
        if line.strip():
            doc.add_paragraph(line.strip())
    
    # 保存文档
    output_path = '/mnt/c/Users/LSD/Desktop/汇报/周报/lsd_20260110_周报.docx'
    
    # 确保目录存在
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    doc.save(output_path)
    print(f"周报已生成: {output_path}")
    
    return output_path

if __name__ == '__main__':
    create_weekly_report()
